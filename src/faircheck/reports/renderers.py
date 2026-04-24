"""Report renderers — PDF (WeasyPrint), Markdown, DOCX (python-docx)."""

from __future__ import annotations

import logging
from pathlib import Path
from typing import Any

from faircheck.reports.charts import generate_charts
from faircheck.reports.data import ReportData
from faircheck.reports.engine import ReportEngine

logger = logging.getLogger(__name__)


class ReportBuilder:
    """High-level API for generating reports in all formats.

    Usage::

        builder = ReportBuilder()
        paths = builder.build(analysis_dict, formats=["pdf", "md", "docx"])
    """

    def __init__(self, engine: ReportEngine | None = None) -> None:
        self.engine = engine or ReportEngine()

    def build(
        self,
        data: dict[str, Any],
        formats: list[str] | None = None,
        output_dir: Path | None = None,
        regulation: str | None = None,
    ) -> dict[str, Path]:
        """Generate reports in requested formats.

        Parameters
        ----------
        data : dict
            Plain dict (from ``AnalysisResult.to_dict()`` or stored session).
        formats : list[str]
            Output formats: ``"pdf"``, ``"md"``, ``"docx"``.  Default: ``["pdf"]``.
        output_dir : Path, optional
            Override output directory.
        regulation : str, optional
            Regulation standard key (e.g. ``"eu_ai_act"``).  When set,
            uses regulation-specific templates with article mappings.

        Returns
        -------
        dict[str, Path]
            Mapping format → output file path.
        """
        formats = formats or ["pdf"]
        report_data = ReportData.from_dict(data)

        # Generate charts and attach to report data
        if report_data.analysis_results:
            report_data.chart_images = generate_charts(report_data.analysis_results)

        out_dir = output_dir or self.engine.get_output_dir(
            report_data.session_id or "default"
        )

        results: dict[str, Path] = {}
        for fmt in formats:
            if fmt == "pdf":
                results["pdf"] = self._render_pdf(report_data, out_dir, regulation)
            elif fmt == "md":
                results["md"] = self._render_md(report_data, out_dir, regulation)
            elif fmt == "docx":
                results["docx"] = self._render_docx(report_data, out_dir)
            else:
                logger.warning("Unknown format '%s', skipping", fmt)

        return results

    def build_html_accessible(
        self,
        data: dict,
        regulation: str | None = None,
    ) -> str:
        """Generate a WCAG AA accessible HTML report (inline string)."""
        report_data = ReportData.from_dict(data)
        if report_data.analysis_results:
            report_data.chart_images = generate_charts(report_data.analysis_results)
        # Attach extra fields for richer template
        report_data._extra = data  # type: ignore[attr-defined]
        return self.engine.render_accessible_html(report_data, regulation=regulation)

    def _render_pdf(
        self, data: ReportData, out_dir: Path, regulation: str | None = None
    ) -> Path:
        """Render HTML → PDF via WeasyPrint (RPT-09)."""
        from weasyprint import HTML

        html_content = self.engine.render_html(data, regulation=regulation)
        out_path = out_dir / "faircheck-report.pdf"
        HTML(string=html_content).write_pdf(str(out_path))
        logger.info("PDF report saved: %s", out_path)
        return out_path

    def _render_md(
        self, data: ReportData, out_dir: Path, regulation: str | None = None
    ) -> Path:
        """Render Markdown report (RPT-10)."""
        md_content = self.engine.render_markdown(data, regulation=regulation)
        out_path = out_dir / "faircheck-report.md"
        out_path.write_text(md_content, encoding="utf-8")
        logger.info("Markdown report saved: %s", out_path)
        return out_path

    def _render_docx(self, data: ReportData, out_dir: Path) -> Path:
        """Render DOCX report via python-docx (RPT-11)."""
        from docx import Document

        doc = Document()

        # Cover (RPT-01)
        doc.add_heading("AI Bias Audit Report", level=0)
        doc.add_paragraph(f"Model: {data.model_name} (v{data.model_version})")
        doc.add_paragraph(f"Audit Date: {data.audit_date}")
        if data.reviewer:
            doc.add_paragraph(f"Reviewer: {data.reviewer}")
        doc.add_paragraph(f"Regulation: {data.regulation_standard}")

        # Executive Summary (RPT-02)
        doc.add_heading("Executive Summary", level=1)
        risk = data.analysis_results.get("overall_risk_level", "unknown")
        doc.add_paragraph(f"Overall Risk Level: {risk.upper()}")

        # Model Card (RPT-03)
        doc.add_heading("Model Card", level=1)
        doc.add_paragraph(
            data.training_data_description or "Not provided."
        )
        doc.add_paragraph(
            f"Intended Use: {data.intended_use or 'Not specified.'}"
        )
        doc.add_paragraph(
            f"Known Limitations: {data.known_limitations or 'None documented.'}"
        )

        # Metrics Table (RPT-04)
        doc.add_heading("Bias Analysis Results", level=1)
        results = data.analysis_results.get("results", {})
        for attr_name, attr_data in results.items():
            doc.add_heading(f"Protected Attribute: {attr_name}", level=2)
            metrics = attr_data.get("metrics", {})
            if metrics:
                table = doc.add_table(rows=1, cols=4)
                table.style = "Table Grid"
                hdr = table.rows[0].cells
                hdr[0].text = "Metric"
                hdr[1].text = "Threshold"
                hdr[2].text = "Value"
                hdr[3].text = "Status"
                for m_name, m_data in metrics.items():
                    row = table.add_row().cells
                    row[0].text = m_name.replace("_", " ").title()
                    row[1].text = str(m_data.get("threshold", ""))
                    row[2].text = (
                        f"{m_data['value']:.4f}"
                        if m_data.get("value") is not None
                        else "N/A"
                    )
                    row[3].text = m_data.get("status", "").upper()

        # Mitigation (RPT-06, D-06)
        doc.add_heading("Mitigation Applied", level=1)
        if data.mitigation:
            doc.add_paragraph(
                f"Algorithm: {data.mitigation.get('algorithm', 'N/A')}"
            )
        else:
            doc.add_paragraph(
                "No mitigation has been applied to this model."
            )

        # Oversight (RPT-07, D-06)
        doc.add_heading("Human Oversight", level=1)
        if data.oversight:
            doc.add_paragraph(
                f"Reviewer: {data.oversight.get('reviewer', '')}"
            )
            doc.add_paragraph(
                f"Decision: {data.oversight.get('decision', '')}"
            )
        else:
            doc.add_paragraph(
                "No human oversight decision has been recorded."
            )

        # Appendix (RPT-08)
        doc.add_heading("Appendix: Raw Data & Methodology", level=1)
        doc.add_paragraph("Metrics computed using FairCheck v0.1.0")
        doc.add_paragraph("Libraries: Fairlearn, AIF360, scikit-learn")

        out_path = out_dir / "faircheck-report.docx"
        doc.save(str(out_path))
        logger.info("DOCX report saved: %s", out_path)
        return out_path
